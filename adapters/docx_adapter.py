import asyncio
import re
import logging
from logging import Logger
from typing import List
from asyncio.events import AbstractEventLoop

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .base_adapter import BaseAdapter

logger: Logger = logging.getLogger(__name__)


def _read_docx_file(file_path: str) -> str:
    try:
        doc: DocxDocument = DocxDocument(file_path)
        
        full_text: List[str] = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text: List[str] = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[HEADER] {para.text}")
            
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[FOOTER] {para.text}")
        
        text: str = "\n".join(full_text)
        
        if not text.strip():
            raise ValueError(
                f"DOCX não contém texto extraível: {file_path}. "
                "Arquivo pode estar vazio ou corrompido."
            )
        
        word_count: int = len(text.split())
        if word_count < 10:
            logger.warning(
                f"DOCX contém muito pouco texto ({word_count} palavras). "
                "Pode estar corrompido ou vazio."
            )
        
        logger.debug(f"Extraído {len(text)} caracteres ({word_count} palavras) do DOCX")
        return text
        
    except PackageNotFoundError:
        raise ValueError(
            f"DOCX corrompido ou inválido: {file_path}. "
            "Arquivo não é um DOCX válido."
        )
    except FileNotFoundError:
        raise IOError(f"Arquivo não encontrado: {file_path}")
    except PermissionError:
        raise IOError(f"Sem permissão para ler o arquivo: {file_path}")
    except Exception as e:
        error_msg: str = str(e).lower()
        if 'password' in error_msg or 'protected' in error_msg or 'encrypted' in error_msg:
            raise ValueError(
                f"DOCX protegido por senha: {file_path}. "
                "Não é possível extrair texto sem a senha."
            )
        raise IOError(f"Erro ao ler o arquivo DOCX '{file_path}': {e}")


class DocxAdapter(BaseAdapter):

    def chunk(self, text: str) -> List[str]:
        paragraphs: List[str] = re.split(r"\n+", text)
        paragraphs: List[str] = [p.strip() for p in paragraphs if p.strip()]
        return paragraphs

    async def read_text(self, file_path: str) -> str:
        loop: AbstractEventLoop = asyncio.get_running_loop()

        text: str = await loop.run_in_executor(None, _read_docx_file, file_path)
        return text


if __name__ == "__main__":
    adapter: DocxAdapter = DocxAdapter()
    print("Adaptador DOCX pronto para ser usado.")
